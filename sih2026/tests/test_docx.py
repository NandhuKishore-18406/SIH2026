"""Tests for DOCX text extraction."""

import docx
import pytest
from sih2026.extraction.docx import extract_docx
from sih2026.extraction.exceptions import (
    NoTextExtractedError,
    UnsupportedFileTypeError,
)


@pytest.fixture
def sample_docx(tmp_path):
    docx_path = tmp_path / "sample_syllabus.docx"
    doc = docx.Document()
    doc.add_heading("Course Syllabus", level=1)
    doc.add_paragraph("UNIT I: Software Engineering Basics")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Module"
    table.cell(0, 1).text = "Hours"
    table.cell(1, 0).text = "Introduction"
    table.cell(1, 1).text = "10"

    doc.save(docx_path)
    return docx_path


@pytest.fixture
def empty_docx(tmp_path):
    docx_path = tmp_path / "empty.docx"
    doc = docx.Document()
    doc.save(docx_path)
    return docx_path


def test_extract_docx_success(sample_docx):
    doc = extract_docx(sample_docx)

    assert doc.filename == "sample_syllabus.docx"
    assert len(doc.pages) == 1
    page = doc.pages[0]

    assert page.page_number == 1
    assert page.extraction_method == "docx"
    assert "Course Syllabus" in page.text
    assert "UNIT I: Software Engineering Basics" in page.text
    assert "Module | Hours" in page.text
    assert "Introduction | 10" in page.text


def test_extract_docx_missing_file():
    with pytest.raises(FileNotFoundError):
        extract_docx("missing.docx")


def test_extract_docx_unsupported_extension(tmp_path):
    invalid_path = tmp_path / "document.pdf"
    invalid_path.write_text("not a docx")
    with pytest.raises(UnsupportedFileTypeError):
        extract_docx(invalid_path)


def test_extract_docx_no_text(empty_docx):
    with pytest.raises(NoTextExtractedError):
        extract_docx(empty_docx)
